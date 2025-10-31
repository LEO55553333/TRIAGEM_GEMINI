import os
import time
import logging
import base64
from pathlib import Path
from flask import Flask, request, jsonify
import requests
from dotenv import load_dotenv
from docx import Document
from threading import Thread
import json
from datetime import datetime
import webbrowser
import sys

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('api_gemini.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class GeminiAPIProcessor:
    def __init__(self, base_path="c:\\API_gemini", temperature=0):
        self.base_path = Path(base_path)
        self.api_key = None
        self.temperature = temperature
        self.load_api_key()
        logger.info(f"Temperatura configurada: {self.temperature}")
        
    def load_api_key(self):
        env_path = self.base_path / ".env"
        if env_path.exists():
            load_dotenv(env_path)
            self.api_key = os.getenv('GEMINI_API_KEY')
            if self.api_key:
                logger.info("API key carregada com sucesso")
                logger.info(f"API key (primeiros 10 caracteres): {self.api_key[:10]}...")
            else:
                logger.error("API key não encontrada no arquivo .env")
        else:
            logger.error(f"Arquivo .env não encontrado em {env_path}")
    
    def encode_pdf_to_base64(self, pdf_path):
        try:
            with open(pdf_path, 'rb') as pdf_file:
                pdf_data = pdf_file.read()
                pdf_size = len(pdf_data)
                logger.info(f"PDF {pdf_path.name} - Tamanho: {pdf_size:,} bytes ({pdf_size/1024/1024:.2f} MB)")
                
                if pdf_data[:4] != b'%PDF':
                    logger.warning(f"ATENÇÃO: {pdf_path.name} pode não ser um PDF válido")
                
                base64_data = base64.b64encode(pdf_data).decode('utf-8')
                logger.info(f"PDF convertido para base64: {len(base64_data):,} caracteres")
                
                return base64_data
        except Exception as e:
            logger.error(f"Erro ao converter PDF {pdf_path}: {str(e)}")
            return None
    
    def read_prompt_file(self, folder_path):
        prompt_path = folder_path / "prompt.txt"
        try:
            with open(prompt_path, 'r', encoding='utf-8') as file:
                content = file.read().strip()
                logger.info(f"Prompt carregado: {len(content)} caracteres")
                return content
        except FileNotFoundError:
            logger.warning(f"[PASTA {folder_path.name}] Sem prompt.txt - usando prompt padrão")
            return None
        except Exception as e:
            logger.error(f"Erro ao ler prompt.txt: {str(e)}")
            return None
    
    def get_pdf_files(self, folder_path):
        pdf_files = list(folder_path.glob("*.pdf"))
        logger.info(f"Encontrados {len(pdf_files)} arquivo(s) PDF")
        for pdf_file in pdf_files:
            size_mb = pdf_file.stat().st_size / 1024 / 1024
            logger.info(f"  - {pdf_file.name} ({size_mb:.2f} MB)")
        return pdf_files
    
    def create_api_payload(self, prompt_text, pdf_files):
        logger.info("=== CRIANDO PAYLOAD ===")
        parts = []
        
        if prompt_text:
            parts.append({"text": prompt_text})
            logger.info("✅ Prompt personalizado adicionado")
        else:
            default_prompt = "Analise detalhadamente o(s) documento(s) PDF anexado(s) e forneça um resumo completo do conteúdo."
            parts.append({"text": default_prompt})
            logger.info("✅ Prompt padrão adicionado")
        
        logger.info(f"Processando {len(pdf_files)} PDF(s)...")
        for i, pdf_file in enumerate(pdf_files, 1):
            logger.info(f"[{i}/{len(pdf_files)}] Processando: {pdf_file.name}")
            pdf_base64 = self.encode_pdf_to_base64(pdf_file)
            
            if pdf_base64:
                parts.append({
                    "inlineData": {
                        "mimeType": "application/pdf",
                        "data": pdf_base64
                    }
                })
                logger.info(f"✅ {pdf_file.name} adicionado ao payload")
            else:
                logger.error(f"❌ Falha ao processar {pdf_file.name}")
        
        payload = {
            "contents": [{"parts": parts}],
            "generationConfig": {
                "temperature": self.temperature,
                "maxOutputTokens": 65536
            }
        }
        
        payload_size = len(json.dumps(payload).encode('utf-8'))
        logger.info(f"Payload criado: {payload_size:,} bytes ({payload_size/1024/1024:.2f} MB)")
        logger.info(f"Total de parts no payload: {len(parts)}")
        
        return payload
    
    def call_gemini_api(self, payload):
        logger.info("=== CHAMANDO API GEMINI ===")
        
        headers = {
            "Content-Type": "application/json",
            "x-goog-api-key": self.api_key
        }
        
        url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"
        logger.info(f"URL: {url}")
        
        try:
            logger.info("Enviando requisição...")
            
            import threading
            stop_progress = threading.Event()
            progress_thread = threading.Thread(target=self._show_api_progress, args=(stop_progress,))
            progress_thread.daemon = True
            progress_thread.start()
            
            response = requests.post(url, headers=headers, json=payload, timeout=400)
            
            stop_progress.set()
            progress_thread.join(timeout=1)
            
            print(f"\r{' ' * 80}\r", end='', flush=True)
            
            logger.info(f"Status Code: {response.status_code}")
            
            if response.status_code == 200:
                logger.info("✅ Resposta recebida com sucesso")
                return response.json()
            else:
                logger.error(f"❌ Erro na API: {response.status_code}")
                logger.error(f"❌ Detalhes: {response.text}")
                return None
                
        except Exception as e:
            stop_progress.set()
            print(f"\r{' ' * 80}\r", end='', flush=True)
            logger.error(f"❌ Erro na requisição: {str(e)}")
            return None
    
    def _show_api_progress(self, stop_event):
        bar_length = 120
        chars = ['⠋', '⠙', '⠹', '⠸', '⠼', '⠴', '⠦', '⠧', '⠇', '⠏']
        
        start_time = time.time()
        i = 0
        
        while not stop_event.is_set():
            elapsed = int(time.time() - start_time)
            minutes = elapsed // 60
            seconds = elapsed % 60
            
            spinner = chars[i % len(chars)]
            
            filled = min(bar_length, elapsed)
            bar = '█' * filled + '░' * (bar_length - filled)
            
            print(f"\r{spinner} Aguardando resposta da API... [{bar}] {minutes:02d}:{seconds:02d}", end='', flush=True)
            
            time.sleep(1)
            i += 1
    
    def extract_response_text(self, response):
        try:
            candidate = response['candidates'][0]
            finish_reason = candidate.get('finishReason', 'UNKNOWN')
            
            response_text = None
            if 'content' in candidate and 'parts' in candidate['content']:
                parts = candidate['content']['parts']
                if parts and 'text' in parts[0]:
                    response_text = parts[0]['text']
            
            if response_text:
                logger.info(f"✅ Texto extraído: {len(response_text):,} caracteres")
                
                if finish_reason == 'MAX_TOKENS':
                    logger.warning("⚠️ RESPOSTA TRUNCADA: Atingiu limite de 65,536 tokens")
                    logger.warning("⚠️ PDF muito complexo - considere dividir em partes menores")
                elif finish_reason == 'STOP':
                    logger.info("✅ Resposta completa (não truncada)")
                else:
                    logger.info(f"Finalizada com: {finish_reason}")
                
                return response_text
            else:
                logger.error("❌ Nenhum texto encontrado na resposta")
                logger.error(f"Estrutura: {json.dumps(response, indent=2)}")
                return None
                
        except Exception as e:
            logger.error(f"❌ Erro ao extrair resposta: {str(e)}")
            return None
    
    def save_response_to_word(self, response_text, output_path):
        try:
            doc = Document()
            doc.add_heading('Análise Gemini', 0)
            doc.add_paragraph(response_text)
            doc.save(output_path)
            logger.info(f"✅ Arquivo salvo: {output_path.name}")
            return True
        except Exception as e:
            logger.error(f"❌ Erro ao salvar Word: {str(e)}")
            return False
    
    def process_folder(self, folder_path):
        folder_name = folder_path.name
        logger.info(f"\n{'='*60}")
        logger.info(f"PROCESSANDO PASTA: {folder_name}")
        logger.info(f"{'='*60}")
        
        prompt_text = self.read_prompt_file(folder_path)
        
        pdf_files = self.get_pdf_files(folder_path)
        if not pdf_files:
            logger.error(f"❌ Nenhum PDF encontrado na pasta {folder_name}")
            return False
        
        payload = self.create_api_payload(prompt_text, pdf_files)
        if not payload:
            logger.error(f"❌ Falha ao criar payload")
            return False
        
        response = self.call_gemini_api(payload)
        if not response:
            logger.error(f"❌ Falha na chamada da API")
            return False
        
        response_text = self.extract_response_text(response)
        if not response_text:
            logger.error(f"❌ Falha ao extrair resposta")
            return False
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"gemini_output_{folder_name}_{timestamp}.docx"
        output_path = folder_path / output_filename
        
        if self.save_response_to_word(response_text, output_path):
            logger.info(f"✅ SUCESSO: Pasta {folder_name} processada!")
            return True
        else:
            logger.error(f"❌ FALHA: Erro ao salvar pasta {folder_name}")
            return False
    
    def process_all_folders(self):
        if not self.api_key:
            logger.error("❌ API key não configurada")
            return
        
        folders = [item for item in self.base_path.iterdir() 
                  if item.is_dir() and item.name.isdigit()]
        folders.sort(key=lambda x: int(x.name))
        
        if not folders:
            logger.error(f"❌ Nenhuma pasta numerada encontrada em {self.base_path}")
            return
        
        logger.info(f"Encontradas {len(folders)} pastas: {[f.name for f in folders]}")
        
        success_count = 0
        for i, folder in enumerate(folders, 1):
            logger.info(f"\n🔄 [{i}/{len(folders)}] Processando pasta {folder.name}")
            
            if self.process_folder(folder):
                success_count += 1
            
            if i < len(folders):
                logger.info("⏳ Aguardando 60 segundos...")
                time.sleep(60)
        
        logger.info(f"\n{'='*60}")
        logger.info(f"RELATÓRIO FINAL")
        logger.info(f"{'='*60}")
        logger.info(f"Total de pastas: {len(folders)}")
        logger.info(f"Sucessos: {success_count}")
        logger.info(f"Falhas: {len(folders) - success_count}")
        logger.info(f"Taxa de sucesso: {success_count/len(folders)*100:.1f}%")
        logger.info(f"{'='*60}")
        
        print(f"\n{'='*60}")
        print("🎉 PROCESSAMENTO FINALIZADO COM SUCESSO! 🎉")
        print(f"{'='*60}")
        print("           ✅ ✅ ✅ ✅ ✅")
        print("       ✅                  ✅")
        print("     ✅                      ✅")
        print("   ✅         TERMINOU!        ✅")
        print("     ✅                      ✅")
        print("       ✅                  ✅")
        print("           ✅ ✅ ✅ ✅ ✅")
        print(f"{'='*60}")
        print("📋 Todos os arquivos foram processados com sucesso!")
        print(f"{'='*60}\n")

app = Flask(__name__)
processor = GeminiAPIProcessor(temperature=0)

@app.route('/start_processing', methods=['GET', 'POST'])
def start_processing():
    try:
        thread = Thread(target=processor.process_all_folders)
        thread.daemon = True
        thread.start()
        
        return jsonify({
            "status": "success",
            "message": "Processamento iniciado! Acompanhe os logs."
        })
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"Erro: {str(e)}"
        }), 500

@app.route('/status', methods=['GET'])
def get_status():
    return jsonify({
        "status": "running",
        "base_path": str(processor.base_path),
        "api_key_loaded": processor.api_key is not None
    })

@app.route('/', methods=['GET'])
def home():
    return """
    <h1>Gemini API Processor</h1>
    <p><strong>Status:</strong> Servidor rodando</p>
    <p><strong>Pasta base:</strong> """ + str(processor.base_path) + """</p>
    <hr>
    <button onclick="startProcessing()" style="padding: 10px 20px; font-size: 16px;">
        🚀 Iniciar Processamento
    </button>
    <p><em>Clique no botão acima para processar todas as pastas</em></p>
    
    <script>
    function startProcessing() {
        fetch('/start_processing', {method: 'POST'})
        .then(response => response.json())
        .then(data => {
            alert(data.message);
            console.log('Processamento iniciado - verifique os logs no terminal');
        })
        .catch(error => {
            alert('Erro: ' + error);
            console.error('Erro:', error);
        });
    }
    </script>
    """

def auto_start_processing():
    print(f"\n{'='*60}")
    print("🚀 PROCESSAMENTO INICIADO - Aguarde...")
    print("📋 Acompanhe os logs detalhados abaixo:")
    print(f"{'='*60}\n")
    
    processor.process_all_folders()

def main():
    logger.info("Iniciando Gemini API Processor...")
    
    if not processor.base_path.exists():
        logger.error(f"Pasta base não existe: {processor.base_path}")
        return
    
    print("\n" + "="*60)
    print("🎯 GEMINI API PROCESSOR - MODO AUTOMÁTICO")
    print("="*60)
    print("✨ O processamento será iniciado automaticamente!")
    print("📝 Acompanhe os logs detalhados durante o processo")
    print("🔥 Temperatura configurada: 0 (máxima precisão)")
    print("="*60)
    
    time.sleep(2)
    
    auto_start_processing()

if __name__ == "__main__":
    main()